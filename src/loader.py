"""
loader.py — Загрузка PDF и DOCX документов.

Извлекает текст из файлов с сохранением метаданных:
- filename: имя файла
- page: номер страницы (для PDF) или 1 (для DOCX)
- source: полный путь к файлу

Возвращает список объектов Document (LangChain), каждый — одна страница.
"""
import os
from typing import List

from langchain_core.documents import Document

# PyMuPDF для PDF (импорт как fitz)
import fitz  # PyMuPDF
from docx import Document as DocxDocument

from src.config import DOCS_DIR, UPLOAD_DIR


def load_pdf(path: str) -> List[Document]:
    """Загрузить PDF постранично. Пропускает повреждённые файлы."""
    filename = os.path.basename(path)
    docs = []

    try:
        pdf = fitz.open(path)
    except Exception as e:
        print(f"  [loader] Пропуск {filename}: повреждённый PDF ({e})")
        return []

    for page_num in range(len(pdf)):
        page = pdf[page_num]
        text = page.get_text("text")  # Извлекаем текст со страницы

        if text.strip():  # Пропускаем пустые страницы
            doc = Document(
                page_content=text,
                metadata={
                    "filename": filename,
                    "page": page_num + 1,  # Нумерация с 1
                    "source": path,
                },
            )
            docs.append(doc)

    pdf.close()
    return docs


def load_docx(path: str) -> List[Document]:
    """
    Загрузить DOCX.
    DOCX не имеет страниц, поэтому весь текст — один Document.
    Пропускает повреждённые/нестандартные файлы без падения.
    """
    filename = os.path.basename(path)

    try:
        docx_doc = DocxDocument(path)
    except Exception as e:
        print(f"  [loader] Пропуск {filename}: повреждённый файл ({e})")
        return []

    try:
        paragraphs = [p.text for p in docx_doc.paragraphs if p.text.strip()]
        text = "\n".join(paragraphs)
    except Exception as e:
        print(f"  [loader] Пропуск {filename}: ошибка чтения ({e})")
        return []

    if not text.strip():
        return []

    return [
        Document(
            page_content=text,
            metadata={
                "filename": filename,
                "page": 1,
                "source": path,
            },
        )
    ]


def load_document(path: str) -> List[Document]:
    """
    Загрузить один документ (PDF или DOCX).
    Автоматически определяет формат по расширению.
    """
    ext = os.path.splitext(path)[1].lower()

    if ext == ".pdf":
        return load_pdf(path)
    elif ext == ".docx":
        return load_docx(path)
    else:
        print(f"[loader] Неподдерживаемый формат: {ext} ({path})")
        return []


def load_all_documents() -> List[Document]:
    """
    Загрузить ВСЕ документы из data/generated и data/uploads.
    Используется при первичной индексации.
    """
    all_docs = []

    # Проверяем обе папки
    for directory in [DOCS_DIR, UPLOAD_DIR]:
        if not os.path.exists(directory):
            continue

        for filename in os.listdir(directory):
            ext = os.path.splitext(filename)[1].lower()
            if ext in (".pdf", ".docx"):
                path = os.path.join(directory, filename)
                docs = load_document(path)
                all_docs.extend(docs)
                print(f"  [loader] {filename}: {len(docs)} стр.")

    print(f"  [loader] Всего загружено: {len(all_docs)} страниц")
    return all_docs
